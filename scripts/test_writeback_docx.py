#!/usr/bin/env python3
"""test_writeback_docx.py — writeback_docx.py 的行为规格（TDD）。

零第三方测试框架：assert + 用例计数。fixture 用 python-docx 现构造小文档
（含标题/正文/表格/merged cell/空 cell/重复短语），不依赖任何真实需求文档。
运行: python3 test_writeback_docx.py   （退出码 0=全过）
"""
import os
import shutil
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import writeback_docx as wb  # noqa: E402

from docx import Document  # noqa: E402

CASES = []


def case(fn):
    CASES.append(fn)
    return fn


# ---------------------------------------------------------------- fixture

def make_fixture_docx(path):
    """构造覆盖关键形态的最小 docx：标题/正文/重复短语/映射表/merged/空 cell。"""
    doc = Document()
    doc.add_heading("引言", level=1)
    doc.add_paragraph("本文档描述客户管理系统的国别风险管理要求。")
    doc.add_heading("新增-业务规则【风险等级】映射表", level=2)
    doc.add_paragraph("国别风险等级映射规则如下，审批流程需要设置阈值。")
    t1 = doc.add_table(rows=6, cols=2)
    for r, (a, b) in enumerate([("AAA", "低风险"), ("AA", "低风险"), ("A", "较低风险"),
                                ("BBB", "较低风险"), ("B", "较低风险"), ("C", "中风险")]):
        t1.cell(r, 0).text = a
        t1.cell(r, 1).text = b
    doc.add_heading("系统设置", level=2)
    doc.add_paragraph("国别风险提示开关可配置，阈值由参数控制。")
    t2 = doc.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "机构类型"
    t2.cell(0, 1).text = "审批层级"
    t2.cell(1, 0).text = "总行"
    t2.cell(1, 1).merge(t2.cell(1, 0))  # merged cell（重复出现同一 cell）
    t3 = doc.add_table(rows=2, cols=2)
    t3.cell(0, 0).text = "字段"
    t3.cell(0, 1).text = ""             # 空 cell
    t3.cell(1, 0).text = "币种"
    t3.cell(1, 1).text = "人民币"
    doc.save(path)


def ledger_line(fid, loc, excerpt, imp="P1", rel="阻断", status="打开",
                desc="问题描述", advice="修改建议", close="关闭条件",
                clar="无", ev="已确认"):
    return ("| " + " | ".join([
        fid, "D3", imp, ev, rel, loc, excerpt, desc, "文档内 §1", advice,
        close, clar, status]) + " |")


HEADER = ("| F-ID | 维度 | impact_level | evidence_status | release_effect | 定位 | 原文摘录 | "
          "问题描述与后果 | 证据来源 | 修改建议 | 关闭条件 | 澄清问题 | 状态 |\n"
          "|---|---|---|---|---|---|---|---|---|---|---|---|---|\n")


def write_ledger(path, lines):
    with open(path, "w", encoding="utf-8") as f:
        f.write(HEADER + "\n".join(lines) + "\n")


def setup_workdir(lines):
    d = tempfile.mkdtemp(prefix="wbtest-")
    make_fixture_docx(os.path.join(d, "source.docx"))
    write_ledger(os.path.join(d, "issues-ledger.md"), lines)
    return d


def load_blocks(path=None):
    d = tempfile.mkdtemp(prefix="wbblk-")
    p = os.path.join(d, "f.docx")
    make_fixture_docx(p)
    doc = Document(p)
    return wb.build_blocks(doc)


# ---------------------------------------------------------------- A 解析索引

@case
def test_blocks_count():
    blocks = load_blocks()
    paras = [b for b in blocks if b.kind == "para"]
    trows = [b for b in blocks if b.kind == "trow"]
    # 段落：3 标题 + 3 正文 = 6；表行：6 + 2 + 2 = 10
    assert len(paras) == 6, f"正文段落 {len(paras)} ≠ 6"
    assert len(trows) == 10, f"表格行 {len(trows)} ≠ 10"
    assert [b.idx for b in blocks] == list(range(len(blocks)))


@case
def test_heading_and_sections():
    doc_blocks = load_blocks()
    doc = None
    secs = wb.build_sections(doc_blocks)
    titles = [doc_blocks[s.start].anchor.text if doc_blocks[s.start].anchor else "?" for s in secs]
    assert len(secs) == 3, f"章节 {len(secs)} ≠ 3"
    assert any("映射表" in s.title_norm for s in secs)
    # 章节区间互不重叠且覆盖到下一标题前
    m = next(s for s in secs if "映射表" in s.title_norm)
    inner = [doc_blocks[i] for i in range(m.start, m.end)]
    assert any(b.kind == "trow" and b.cells_norm == ["AAA", "低风险"] for b in inner)


@case
def test_table_merged_and_empty_cells():
    blocks = load_blocks()
    t2_rows = [b for b in blocks if b.kind == "trow" and "总行" in b.cells_norm]
    assert t2_rows, "merged 表行未入索引"
    assert t2_rows[0].cells_norm == ["总行"], f"merged 去重失败: {t2_rows[0].cells_norm}"
    t3_rows = [b for b in blocks if b.kind == "trow" and "字段" in b.cells_norm]
    assert t3_rows[0].cells_norm == ["字段"], f"空 cell 应被剔除: {t3_rows[0].cells_norm}"
    assert t3_rows[0].anchor is not None  # 行内有非空 cell 可锚


@case
def test_empty_paragraph_unanchorable():
    blocks = load_blocks()
    # fixture 无空段，构造性检查：heading 块 anchor 存在、所有 trow anchor 非 None 或行为已知
    for b in blocks:
        if b.kind == "para" and not b.text_norm:
            assert b.anchor is None


# ---------------------------------------------------------------- B 摘录拆分

@case
def test_split_plus_and_vs():
    # 引号在拆分时剥掉（docx 原文无包裹引号，剥掉利于匹配）
    assert wb.split_excerpts('"有效期1年" + "有效期为2年"') == ["有效期1年", "有效期为2年"]
    assert len(wb.split_excerpts("A vs B")) == 2
    assert wb.split_excerpts('"X"（三处同源：§2.1/§3.2）') == ["X"]


@case
def test_atomize_table():
    a = wb.atomize("| B | 较低风险 |")
    assert a.kind == "table" and a.tokens == ["B", "较低风险"], (a.kind, a.tokens)
    a2 = wb.atomize("|B|较低风险|")
    assert a2.kind == "table" and a2.tokens == ["B", "较低风险"]


@case
def test_atomize_pair_multi_left():
    """pair 左侧斜线多值：AAA/AA→低风险。tokens=([AAA,AA], 低风险)。"""
    a = wb.atomize("AAA/AA→低风险")
    assert a.kind == "pair", a.kind
    lefts, right = a.tokens
    assert sorted(lefts) == ["AA", "AAA"] and right == "低风险", a.tokens


@case
def test_atomize_text():
    a = wb.atomize("审批流程需要设置阈值")
    assert a.kind == "text" and a.tokens == [wb.normalize("审批流程需要设置阈值")]


# ---------------------------------------------------------------- C 归一化

@case
def test_normalize():
    assert wb.normalize("有效期 １年") == "有效期1年"          # 全角+空白
    assert wb.normalize("“引号”文本") == '"引号"文本'
    assert wb.normalize("") == ""


# ---------------------------------------------------------------- D 分层匹配

@case
def test_L1_unique_and_multi():
    blocks = load_blocks()
    cand = set(range(len(blocks)))
    cr = wb.match_chunk("审批流程需要设置阈值", blocks, cand, 10)
    assert cr.layer == "L1" and len(cr.hits) == 1, (cr.layer, cr.hits)
    cr2 = wb.match_chunk("国别风险", blocks, cand, 10)
    assert len(cr2.hits) == 3, f"高频词应 3 处: {cr2.hits}"  # 引言段+规则段+设置段


@case
def test_L15_truncated():
    blocks = load_blocks()
    cand = set(range(len(blocks)))
    cr = wb.match_chunk("审批流程需要设置阈值超出部分截断掉的尾巴", blocks, cand, 10)
    assert cr.layer == "L1.5" and len(cr.hits) == 1, (cr.layer, cr.hits)


@case
def test_L2_table_row():
    blocks = load_blocks()
    cand = set(range(len(blocks)))
    cr = wb.match_chunk("| B | 较低风险 |", blocks, cand, 10)
    assert cr.layer == "L2" and len(cr.hits) == 1, (cr.layer, cr.hits)
    b = blocks[cr.hits[0]]
    assert b.cells_norm == ["B", "较低风险"] and b.anchor is not None
    # 单向包含：'B' 不得子串命中 'BBB' 行（评级 B≠BBB，挂错行即误导）
    cr2 = wb.match_chunk("| BBB | 较低风险 |", blocks, cand, 10)
    assert len(cr2.hits) == 1 and blocks[cr2.hits[0]].cells_norm == ["BBB", "较低风险"], \
        (cr2.layer, cr2.hits)


@case
def test_L25_pair_hits_two_rows():
    """AAA/AA→低风险 应命中 AAA 行与 AA 行两处（多值左 token 语义）。"""
    blocks = load_blocks()
    cand = set(range(len(blocks)))
    cr = wb.match_chunk("AAA/AA→低风险", blocks, cand, 10)
    assert cr.layer == "L2.5" and len(cr.hits) == 2, (cr.layer, cr.hits)
    got = [blocks[i].cells_norm for i in cr.hits]
    assert ["AAA", "低风险"] in got and ["AA", "低风险"] in got, got


@case
def test_L3_lowconf_and_status():
    blocks = load_blocks()
    doc = Document()
    sections = wb.build_sections(blocks)
    row = {"F-ID": "F-001", "维度": "D3", "impact_level": "P2", "evidence_status": "较高可信",
           "release_effect": "条件放行", "定位": "",
           "原文摘录": "国别风险、映射规则、审批阈值",  # 非子串（顿号连接+一处换词）→ 只能 L3 token 覆盖
           "问题描述与后果": "x", "证据来源": "s", "修改建议": "m", "关闭条件": "c",
           "澄清问题": "无", "状态": "打开"}
    fr = wb.match_finding(row, blocks, sections, 10)
    assert fr.status == "低置信", fr.status  # 短语被改写（少了后半句）→ 只能 token 覆盖命中


@case
def test_L4_section_fallback():
    blocks = load_blocks()
    sections = wb.build_sections(blocks)
    row = {"F-ID": "F-002", "定位": "§3.1.1.1 新增-业务规则【风险等级】映射表",
           "原文摘录": "这句话在文档中不存在", "状态": "打开"}
    fr = wb.match_finding(row, blocks, sections, 10)
    assert fr.status == "章节降级", fr.status
    hb = blocks[fr.anchors[0]]
    assert hb.heading and "映射表" in hb.heading


@case
def test_unmatched():
    blocks = load_blocks()
    sections = wb.build_sections(blocks)
    row = {"F-ID": "F-003", "定位": "", "原文摘录": "完全不相关的句子", "状态": "打开"}
    fr = wb.match_finding(row, blocks, sections, 10)
    assert fr.status == "未匹配" and fr.anchors == []


@case
def test_section_filter_narrows():
    """高频词"国别风险"全文 3 处；定位限定章节后只剩区间内命中。"""
    blocks = load_blocks()
    sections = wb.build_sections(blocks)
    row = {"F-ID": "F-004", "定位": "§3.1.1.1 新增-业务规则【风险等级】映射表",
           "原文摘录": "国别风险", "状态": "打开"}
    fr = wb.match_finding(row, blocks, sections, 10)
    assert fr.status == "已批注", fr.status
    assert len(fr.anchors) == 1, f"节内应只 1 处'国别风险': {fr.anchors}"


@case
def test_cap():
    blocks = load_blocks()
    cand = set(range(len(blocks)))
    cr = wb.match_chunk("国别风险", blocks, cand, 2)
    assert len(cr.hits) == 2 and cr.total_hits == 3 and cr.capped is True


# ---------------------------------------------------------------- E 批注文本

@case
def test_comment_text_full():
    row = {"F-ID": "F-001", "impact_level": "P1", "release_effect": "阻断（条件：修订）",
           "evidence_status": "已确认", "问题描述与后果": "与 F-009 矛盾，§3.2 口径不一",
           "修改建议": "统一口径", "关闭条件": "两处一致", "澄清问题": "以哪个为准？",
           "定位": "§2.1"}
    t = wb.build_comment_text(row, seq_total=3, seq_i=2, table_rows=None, low_conf=False)
    assert "【评审意见 F-001】P1 · 阻断（已确认）" in t
    assert "问题：与 F-009 矛盾，第3.2节 口径不一" in t   # §3.2 → 第3.2节
    assert "关闭条件：两处一致" in t
    assert "待澄清：以哪个为准？" in t
    assert "关联：F-009" in t
    assert "本问题共 3 处标注（本处 2/3）" in t
    assert t.endswith(wb.SIGNATURE)


@case
def test_comment_text_p2_and_table_merge():
    row = {"F-ID": "F-002", "impact_level": "P2", "release_effect": "不影响放行",
           "evidence_status": "较高可信", "问题描述与后果": "弱词", "修改建议": "改明确",
           "关闭条件": "c", "澄清问题": "无", "定位": ""}
    t = wb.build_comment_text(row, seq_total=1, seq_i=1, table_rows=5, low_conf=False)
    assert "关闭条件" not in t          # P2 不带关闭条件
    assert "待澄清" not in t            # 澄清=无 省行
    assert "本问题涉及本表 5 行" in t
    assert "共 1 处" not in t           # table_rows 优先，不再出共 N 处


@case
def test_fmt_text():
    assert wb.fmt_text("A \\| B") == "A | B"
    assert wb.fmt_text("见 §3.1.1.2 及 §2") == "见 第3.1.1.2节 及 第2节"
    assert wb.fmt_text("**加粗**`代码`") == "加粗代码"


@case
def test_plan_comment_points_merges_table():
    blocks = load_blocks()
    t1_id = next(b.table_id for b in blocks if b.cells_norm == ["AAA", "低风险"])
    same_table = [b.idx for b in blocks if b.table_id == t1_id][:3]
    pts = wb.plan_comment_points(same_table, blocks)
    assert len(pts) == 1 and pts[0][1] == 3, pts  # 3 行合并 1 条批注
    mixed = [same_table[0], next(b.idx for b in blocks if b.kind == "para" and b.text_norm and "阈值" in b.text_norm)]
    pts2 = wb.plan_comment_points(mixed, blocks)
    assert len(pts2) == 2, pts2


# ---------------------------------------------------------------- F 端到端

@case
def test_run_full_pipeline():
    d = setup_workdir([
        ledger_line("F-001", "§3.1.1.1 新增-业务规则【风险等级】映射表", "\\| B \\| 较低风险 \\|",
                    desc="映射顺序错乱（关联 F-002）"),
        ledger_line("F-002", "§2.1", "审批流程需要设置阈值", imp="P2", rel="条件放行"),
    ])
    try:
        rc = wb.run(d)
        assert rc == 0, f"退出码 {rc} ≠ 0"
        out = [f for f in os.listdir(d) if "评审批注" in f]
        assert len(out) == 1, out
        # 状态文件
        st = open(os.path.join(d, "docx-writeback.md"), encoding="utf-8").read()
        assert "| F-001 | 1 | 1 | 已批注" in st, st
        assert "| F-002 | 1 | 1 | 已批注" in st
        assert "覆盖率: 2/2" in st
        # 读回：RR 批注数 = 状态文件已标注处数和
        n = wb.count_rr_comments(os.path.join(d, out[0]))
        assert n == 2, f"RR 批注 {n} ≠ 2"
    finally:
        shutil.rmtree(d)


@case
def test_run_unmatched_exit_code():
    lines = [
        ledger_line("F-001", "§2.1", "审批流程需要设置阈值"),
        ledger_line("F-002", "", "根本不存在的摘录内容_xyz"),
    ]
    d = setup_workdir(lines)
    try:
        rc = wb.run(d)
        assert rc == 1, f"有未匹配应退出 1，实际 {rc}"
        rc2 = wb.run(d, allow_unmatched=True)
        assert rc2 == 0, f"--allow-unmatched 应退出 0，实际 {rc2}"
        st = open(os.path.join(d, "docx-writeback.md"), encoding="utf-8").read()
        assert "| F-002 | 1 | 0 | 未匹配" in st and "未匹配明细" in st
        assert "--allow-unmatched" in st
    finally:
        shutil.rmtree(d)


@case
def test_source_sha256_untouched():
    import hashlib
    d = setup_workdir([ledger_line("F-001", "§2.1", "审批流程需要设置阈值")])
    try:
        p = os.path.join(d, "source.docx")
        h0 = hashlib.sha256(open(p, "rb").read()).hexdigest()
        wb.run(d)
        h1 = hashlib.sha256(open(p, "rb").read()).hexdigest()
        assert h0 == h1, "source.docx 被改动（违反只读铁律）"
    finally:
        shutil.rmtree(d)


@case
def test_preexisting_comments_preserved_and_filtered():
    d = setup_workdir([ledger_line("F-001", "§2.1", "审批流程需要设置阈值")])
    try:
        # 预挂一条他人批注（模拟 CR261622 的 mayinyin 场景）
        p = os.path.join(d, "source.docx")
        doc = Document(p)
        paras = [x for x in doc.paragraphs if "国别风险管理要求" in x.text]
        doc.add_comment(runs=paras[0].runs, text="旧批注", author="someone-else", initials="SE")
        doc.save(p)
        rc = wb.run(d)
        assert rc == 0
        out = [f for f in os.listdir(d) if "评审批注" in f][0]
        doc2 = Document(os.path.join(d, out))
        authors = [c.author for c in doc2.comments]
        assert authors.count("someone-else") == 1, f"原有批注丢失: {authors}"
        assert authors.count(wb.AUTHOR) == 1, f"RR 批注数被污染: {authors}"
        assert wb.count_rr_comments(os.path.join(d, out)) == 1
    finally:
        shutil.rmtree(d)


# ---------------------------------------------------------------- main

def main():
    failed = []
    for fn in CASES:
        try:
            fn()
            print(f"  PASS  {fn.__name__}")
        except AssertionError as e:
            failed.append((fn.__name__, str(e)))
            print(f"  FAIL  {fn.__name__}: {e}")
        except Exception as e:  # noqa: BLE001
            failed.append((fn.__name__, f"{type(e).__name__}: {e}"))
            print(f"  ERROR {fn.__name__}: {type(e).__name__}: {e}")
    print(f"\n{len(CASES) - len(failed)}/{len(CASES)} 通过")
    if failed:
        sys.exit(1)


if __name__ == "__main__":
    main()
