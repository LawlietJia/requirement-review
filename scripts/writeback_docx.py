#!/usr/bin/env python3
"""writeback_docx.py — 评审台账回写 docx 批注副本（requirement-review V1.7.0）。

读 <工作目录>/source*.docx + issues-ledger.md，生成两件套：
  <原文件名>-评审批注-<YYYYMMDD>.docx  批注副本（每个 Finding 的 Word 批注直挂原文处）
  docx-writeback.md                    回写状态清单（validate_review.py section 11 校验对象）

铁律：
  - 原件只读不动（打开读内存，另存副本；运行前后 sha256 比对机检）
  - 原件自带他人批注保留不动；所有批注计数按 author="requirement-review" 过滤
  - 批注文本纯组装零改写——通俗化由台账源头（report-template.md 措辞规范）保证
  - 匹配失败的 Finding 不得静默丢失：进未匹配明细，退出码 1（--allow-unmatched 豁免）

用法: python3 writeback_docx.py <评审工作目录> [--allow-unmatched] [--cap N]
退出码: 0=全部 Finding 至少 1 锚（含章节降级/低置信）；1=存在零锚 Finding；2=输入错误
依赖: python-docx>=1.2.0（add_comment）
"""
from __future__ import annotations

import argparse
import datetime
import hashlib
import json
import os
import re
import sys
import unicodedata
import zipfile
from dataclasses import dataclass, field

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from validate_review import parse_ledger  # noqa: E402

try:
    import docx as _docx_pkg
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("错误: 需要 python-docx>=1.2.0（uv pip install python-docx）", file=sys.stderr)
    sys.exit(2)

AUTHOR = "requirement-review"
INITIALS = "RR"
SIGNATURE = "—— requirement-review 自动回写"
DEFAULT_CAP = 3   # 多处命中的挂注上限。曾为 10——通用短句（"参见《总册》"）会撒网近全文，
                  # 读者看到的批注 2/3 与该 Finding 无关（真实银行软需实测）；真需逐处提示的
                  # （如功能编号全占位）由批注文本写明"全文共 N 处"，不必逐处挂注。
L3_THRESHOLD = 0.6
VALID_WB_STATUS = ("已批注", "低置信", "章节降级", "未匹配")
COMMENT_COUNT_PAT = re.compile(r'<w:comment [^>]*?w:id="(\d+)"[^>]*?w:author="([^"]*)"')


# ---------------------------------------------------------------- 归一化

def normalize(text: str) -> str:
    """NFKC + 去零宽 + 统一引号 + 折叠全部空白。匹配两端同归一化，只增容忍不改语义。"""
    if not text:
        return ""
    t = unicodedata.normalize("NFKC", str(text))
    for zw in ("​", "‌", "‍", "﻿", " "):
        t = t.replace(zw, "")
    t = re.sub(r"\\([*_\[\]()#<>~`|])", r"\1", t)  # 剥 md 表格转义（\* 等）——docx 原文无反斜杠，不剥则含公式摘录永不中（真实银行软需实测）
    for a, b in (("“", '"'), ("”", '"'), ("‘", "'"), ("’", "'")):
        t = t.replace(a, b)
    return re.sub(r"\s+", "", t)


# ---------------------------------------------------------------- 数据结构

@dataclass
class Atom:
    """摘录 chunk 的匹配原子。table=竖线行（tokens=cell 归一化值）；pair=A→B 记法；text=普通文本。"""
    kind: str            # "text" | "table" | "pair"
    raw: str
    tokens: list


@dataclass
class Block:
    """文档流块：正文段落或表格行（平铺）。anchor 为可挂批注的段落对象（零 run 时 None）。"""
    idx: int
    kind: str            # "para" | "trow"
    anchor: object
    text_norm: str
    cells_norm: list
    table_id: int        # 段落为 -1
    row_idx: int
    heading: str         # 归一化标题文本，非标题为 ""
    is_toc: bool = False # 目录条目段——永不作锚点（摘录与目录行同文时曾误挂目录）
    level: int = 0       # 标题层级（Heading N→N，非标题 0）——供区间含子节的跨度计算


@dataclass
class Section:
    title_norm: str
    start: int           # 标题块 idx（含）
    end: int             # 下一标题块 idx（不含，原子段——不含子节）
    level: int = 1       # 标题层级

    def span_end(self, sections) -> int:
        """完整跨度终点（**含全部子节**）：到下一个 level<=自身 的标题块前。"""
        for k in range(0, len(sections)):
            if sections[k].start == self.start:
                for s2 in sections[k + 1:]:
                    if s2.level <= self.level:
                        return s2.start
                return sections[-1].end if sections else self.end
        return self.end


@dataclass
class ChunkResult:
    chunk: str
    hits: list           # 命中块 idx（已 cap）
    layer: str | None    # L1/L1.5/L2/L2.5/L3/None
    total_hits: int      # cap 前命中总数
    capped: bool


@dataclass
class FindingResult:
    fid: str
    row: dict
    chunks: list         # ChunkResult
    anchors: list        # 合并后的批注挂点块 idx（有序去重）
    status: str          # 已批注/低置信/章节降级/未匹配
    notes: list
    placed: int = 0      # 实际写入批注处数（write_comments 时回填）


# ---------------------------------------------------------------- A 解析索引

def split_excerpts(excerpt: str) -> list:
    """摘录 → chunk 列表（多处证据）。按 ` + ` / ` vs ` 拆，剥同源注记括号。"""
    e = re.sub(r"（[\d一二三四五六七八九十]+\s*处同源[：:][^）]*）", "", excerpt or "")
    parts = re.split(r"\s*\+\s*|\s+vs\s+", e)
    out = [p.strip(" \t\"'“”‘’") for p in parts]
    return [p for p in out if p]


def atomize(chunk: str) -> Atom:
    """chunk → 匹配原子。含竖线→table；含 →/= →pair；否则 text。"""
    c = chunk.replace("\\|", "|")
    if "|" in c:
        cells = [normalize(x) for x in c.strip().strip("|").split("|")]
        cells = [x for x in cells if x]
        if len(cells) >= 1:
            return Atom("table", chunk, cells)
    m = re.split(r"[→]|＝|(?<![!<>=])=(?!=)", c, maxsplit=1)
    if len(m) == 2 and normalize(m[0]) and normalize(m[1]):
        # 左侧多值：`AAA/AA→低风险` 表示 AAA 与 AA 都映射到低风险——命中表内多行
        lefts = [normalize(x) for x in re.split(r"[/、]", m[0]) if normalize(x)]
        return Atom("pair", chunk, (lefts, normalize(m[1])))
    return Atom("text", chunk, [normalize(c)])


def is_heading(para) -> bool:
    try:
        return bool(para.style) and (para.style.name or "").startswith("Heading")
    except Exception:
        return False


def heading_level(para) -> int:
    """标题层级：Heading N / 标题 N → N；非标题 0。层级用于区间跨度——
    章节区间必须**包含全部子节**（『国别风险评级查询』到其『主界面』子标题就断了，
    子节里的列表表格全被区间挡住——真实银行软需实测 F-004 万元句因此未定位）。"""
    try:
        name = (para.style.name or "") if para.style else ""
    except Exception:
        return 0
    m = re.search(r"(?:Heading|标题)\s*(\d+)", name)
    return int(m.group(1)) if m else 0


def _full_text(element) -> str:
    """XML 层全量文本：遍历全部 w:t 节点（**含 w:hyperlink 内的 run**）。

    python-docx 的 para.text / cell.text 不含超链接内文本——真实银行软需实测 F-019
    的摘录「点击跳转额度中心"国别风险-国别风险限额台账"页面」整句在 hyperlink 里，
    para.text 只剩 '额度中心""展示限额使用情况。'，按 .text 建索引必然漏检。"""
    from docx.oxml.ns import qn
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def _all_runs(para):
    """段落内全部 Run（**含 w:hyperlink 内的 run**——para.runs 只暴露普通 run，
    纯超链接段落会得到空列表导致批注无处可挂）。"""
    from docx.oxml.ns import qn
    from docx.text.run import Run
    return [Run(r, para) for r in para._p.iter(qn("w:r"))]


def _looks_like_toc(raw_text: str, para_xml: str, rel_pos: float) -> bool:
    """目录条目段识别：①段内含 _Toc 书签超链接（Word 目录条目的结构特征）；
    ②或位于文档前 20% 且为"短文本+行尾页码"形态（如『3.5. 数据管理目标- 62 -』）。
    目录行永不作锚点——摘录取自目录形态文本时，正文命中之前会先撞上目录（真实银行软需实测）。"""
    if "_Toc" in (para_xml or ""):
        return True
    t = (raw_text or "").strip()
    if rel_pos <= 0.2 and 0 < len(t) <= 60 and re.search(r"[-–—]?\s*\d{1,3}\s*[-–—]?\s*$", t):
        return True
    return False


def build_blocks(doc) -> list:
    """按文档流平铺块索引：段落 + 表格行（嵌套表递归，行归并入所属外层表）。"""
    blocks: list[Block] = []
    counter = [0]
    table_seq = [0]

    def walk_table(table, table_id):
        for r_i, row in enumerate(table.rows):
            cells, seen, anchor = [], set(), None
            row_text_parts = []
            try:
                row_cells = row.cells
            except Exception:
                row_cells = []
            for cell in row_cells:
                cn = normalize("".join(_full_text(p._p) for p in cell.paragraphs))
                if cn and cn not in seen:
                    seen.add(cn)
                    cells.append(cn)
                    row_text_parts.append(cn)
                for para in cell.paragraphs:
                    if anchor is None and _all_runs(para) and _full_text(para._p).strip():
                        anchor = para
                for nested in cell.tables:
                    walk_table(nested, table_id)
            if cells:
                blocks.append(Block(counter[0], "trow", anchor,
                                    "".join(row_text_parts), cells, table_id, r_i, ""))
                counter[0] += 1

    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            full = _full_text(item._p)
            lv = heading_level(item)
            h = normalize(full) if lv else ""
            anchor = item if (_all_runs(item) and full.strip()) else None
            toc = _looks_like_toc(full, item._p.xml, counter[0] / max(1, len(doc.paragraphs)))
            blocks.append(Block(counter[0], "para", anchor, normalize(full), [], -1, -1, h, toc, lv))
            counter[0] += 1
        elif isinstance(item, Table):
            walk_table(item, table_seq[0])
            table_seq[0] += 1
    return blocks


def build_sections(blocks) -> list:
    secs, h_idxs = [], [b for b in blocks if b.heading]
    for j, b in enumerate(h_idxs):
        end = h_idxs[j + 1].idx if j + 1 < len(h_idxs) else len(blocks)
        secs.append(Section(b.heading, b.idx, end, b.level or 9))
    return secs


def _loc_words(loc: str) -> tuple:
    """定位字段 → (标题词候选, 章节号列表, §节名列表)。
    - 章节号（§3.5）单独提取：标题带编号的文档可直接前缀匹配；
    - §节名（§3.1.2.1 国别风险外部评级结果-新增 → 『国别风险外部评级结果』）：
      **银行软需 docx 的正文标题通常不带编号**（编号在目录行/自动编号里，正文标题是
      纯文字），§号本身匹配不到标题，必须拿 §号后紧跟的节名对齐标题（真实银行软需实测）。
      节名取 §号后到 -/（/【 边界的第一段，≥3 字（『新增』『详情』等 2 字节名歧义
      过大，放弃——由整串/摘录兜底）。
    - 整串（去编号与圆括号注记）与【】内词仍提取。括号注记（（L1816-1818）（8 节）
      等）必须剥除，否则整串永远匹配不到标题（真实银行软需实测）。"""
    s = loc or ""
    nums = [n.rstrip(".") for n in re.findall(r"[§]\s*([\d.]+)", s)]
    names = []
    for m in re.findall(r"[§]\s*[\d.]+\s*([^\s（(【§+]+)", s):
        first = re.split(r"[-—–]", m)[0]
        n = normalize(first)
        if len(n) >= 3:
            names.append(n)
    s2 = re.sub(r"（[^）]*）", "", s)
    s2 = re.sub(r"[§]\s*[\d.]+\s*", "", s2)
    key = normalize(s2)
    words = [normalize(w) for w in re.findall(r"【([^】]{2,})】", s)]
    out = [k for k in [key] + words if len(k) >= 2]
    return list(dict.fromkeys(out)), nums, list(dict.fromkeys(names))


def _title_starts_with_num(title_norm: str, num: str) -> bool:
    """标题是否以章节号开头（归一化后无空格，如『3.5.数据管理目标』）。
    要求编号后紧跟分隔符，防『3.5』误配『3.50』『3.55』等更长编号。"""
    return (title_norm.startswith(num + ".")
            or title_norm.startswith(num + "、")
            or title_norm.startswith(num + "．")
            or title_norm == num)


def section_candidates(sections, blocks, loc) -> tuple:
    """返回 (候选块 idx 集合, 过滤成功?)。四层，逐层收紧：
    ①章节号前缀（标题带编号的文档）；②§节名（标题 == 节名 或 startswith 节名）；
    ③整串 key（≥4 字，substring in 标题）；④【】词组——每个词都必须**恰等于**某标题
    才取该标题区间（substring 的 any/all 都会把『国家（地区）』错配到『国家（地区）
    风险系数及转换系数』等含同词的无关章节，真实银行软需实测 F-011/F-028 降级锚因此挂错）。
    全失败退全文（filtered=False，由 match_finding 施加『仅挂首处』约束）。"""
    all_idx = set(range(len(blocks)))
    words, nums, names = _loc_words(loc)
    union = set()

    def take(sec):
        union.update(range(sec.start, sec.span_end(sections)))

    if nums:
        for sec in sections:
            if any(_title_starts_with_num(sec.title_norm, n) for n in nums):
                take(sec)
    if not union and names:
        for sec in sections:
            if any(sec.title_norm == n or sec.title_norm.startswith(n) or n.startswith(sec.title_norm)
                   for n in names):
                take(sec)
    if not union and words:
        for w in words:
            if len(w) >= 4:
                for sec in sections:
                    if w in sec.title_norm:
                        take(sec)
    if not union and words:
        heads = {w: [s for s in sections if s.title_norm == w] for w in words}
        if all(hs for hs in heads.values()):
            for hs in heads.values():
                for sec in hs:
                    take(sec)
    return (union if union else all_idx), bool(union)


# ---------------------------------------------------------------- B 分层匹配

def _apply_cap(hits: list, cap: int):
    return hits[:cap], len(hits), len(hits) > cap


def _tokenize(chunk: str) -> list:
    c = chunk.replace("\\|", " ")
    c = re.sub(r"[【】《》\[\]()（）·;；,，.。:：!！?？\-—/\\|\"'“”‘’+＋、\s]+", " ", c)
    return [t for t in c.split() if len(t) >= 2]


def _is_num_token(t: str) -> bool:
    """纯数字类 token（含小数点/百分号），如 30 / 2.00 / 95%。"""
    return bool(re.fullmatch(r"[\d.,]+%?", t))


def _l3_match(chunk: str, blocks, cand, cap: int):
    toks = _tokenize(chunk)
    if not toks:
        return [], 0, False
    # 数字类 token 必须整词边界命中（"30" 不得子串命中 "20260630"——真实银行软需实测
    # 修订记录日期行因此被误挂）；tokenize 滤掉单字符后只剩数字 token 时（表格行摘录
    # "\| 7 \| D \| 0 \| 是 \| 30 \| 是 \|" 只剩 "30"），单个数字的覆盖度完全不可信，
    # 拒绝 L3 兜底，让该 chunk 走未定位路径。
    if all(_is_num_token(t) for t in toks):
        return [], 0, False
    hits = []
    for i in cand:
        bn = blocks[i].text_norm
        cov = 0
        for t in toks:
            if _is_num_token(t):
                if re.search(rf"(?<![\d.]){re.escape(t)}(?![\d.])", bn):
                    cov += 1
            elif t in bn:
                cov += 1
        if cov / len(toks) >= L3_THRESHOLD:
            hits.append(i)
    return _apply_cap(hits, cap)


def _cell_match(nc: str, hc: str) -> bool:
    """摘录 cell(nc) 与块行 cell(hc) 的匹配：默认精确相等；仅当块 cell ≥3 字符且为
    摘录子串时容错（摘录更详细，如摘录带单位/括注）。短码值（A/B/C 等 ≤2 字符）
    一律精确——评级 B 与 BBB 互相子串，任何方向的包含匹配都会挂错行。"""
    if nc == hc:
        return True
    if len(hc) >= 3 and hc in nc:
        return True
    return False


def _subseq_match(needle: list, hay: list) -> bool:
    """needle（摘录 cell 序列）是 hay（块行 cell 序列）的连续子序列（逐 cell _cell_match）。"""
    n, h = len(needle), len(hay)
    if n == 0 or n > h:
        return False
    for s in range(h - n + 1):
        if all(_cell_match(needle[k], hay[s + k]) for k in range(n)):
            return True
    return False


def match_chunk(chunk: str, blocks, cand, cap: int) -> ChunkResult:
    """单 chunk 分层匹配：L1 子串→L1.5 前缀容错→L2 表格行→L2.5 PAIR→L3 token 兜底（低置信）。
    目录条目块（is_toc）全层排除——摘录与目录行同文时（如『3.5. …- 62 -』形态）不得挂目录。"""
    atom = atomize(chunk)
    ordered = [i for i in sorted(cand) if not blocks[i].is_toc]

    if atom.kind == "text":
        n = atom.tokens[0]
        if n:
            hits = [i for i in ordered if n in blocks[i].text_norm]
            if hits:
                h, t, c = _apply_cap(hits, cap)
                return ChunkResult(chunk, h, "L1", t, c)
            if len(n) >= 6:
                # 截断摘录容错：头部是原文子串、尾部有拼接偏差——前缀从长到短递减重试
                for ratio in (0.8, 0.6, 0.45):
                    pref = n[: max(6, int(len(n) * ratio))]
                    hits = [i for i in ordered if pref in blocks[i].text_norm]
                    if hits:
                        h, t, c = _apply_cap(hits, cap)
                        return ChunkResult(chunk, h, "L1.5", t, c)
    elif atom.kind == "table":
        hits = [i for i in ordered
                if blocks[i].kind == "trow" and _subseq_match(atom.tokens, blocks[i].cells_norm)]
        if hits:
            h, t, c = _apply_cap(hits, cap)
            return ChunkResult(chunk, h, "L2", t, c)
        joined = "".join(atom.tokens)
        if joined:
            hits = [i for i in ordered if joined in blocks[i].text_norm]
            if hits:
                h, t, c = _apply_cap(hits, cap)
                return ChunkResult(chunk, h, "L2.5", t, c)
    else:  # pair（左侧多值：任一左 token 与右 token 同块即命中）
        lefts, right = atom.tokens
        hits = [i for i in ordered
                if right in blocks[i].text_norm
                and any(l in blocks[i].text_norm for l in lefts)]
        if hits:
            h, t, c = _apply_cap(hits, cap)
            return ChunkResult(chunk, h, "L2.5", t, c)

    h, t, c = _l3_match(chunk, blocks, ordered, cap)
    if h:
        return ChunkResult(chunk, h, "L3", t, c)
    return ChunkResult(chunk, [], None, 0, False)


def match_finding(row, blocks, sections, cap) -> FindingResult:
    loc = row.get("定位", "")
    cand, filtered = section_candidates(sections, blocks, loc)
    chunks = split_excerpts(row.get("原文摘录", ""))
    cres = [match_chunk(c, blocks, cand, cap) for c in chunks]

    if not filtered:
        # 定位区间未解析（cand 退化为全文）：通用短句全文撒网风险最高（"参见《总册》"
        # 类句子全文十余处，命中前 3 处也多半与该 Finding 无关）——每 chunk 只挂首处，
        # 其余在 note 声明（真实银行软需实测：F-038 曾因此把 1 处问题挂成 10 处）。
        for cr in cres:
            if len(cr.hits) > 1:
                cr.total_hits = max(cr.total_hits, len(cr.hits))
                cr.hits = cr.hits[:1]
                cr.capped = True
    else:
        # 区间已解析但个别 chunk 零命中：长句（归一化 ≥20 字）在全文内基本唯一，
        # 回退全文找 1 处（定位里 vs 双证据跨节时，另一节的 chunk 会被区间挡住——
        # 真实银行软需 F-004 的公式句在『新增』节而区间解析到『查询』节）。
        all_idx = sorted(set(range(len(blocks))))
        for cr in cres:
            if not cr.hits and len(normalize(cr.chunk.replace("\\|", ""))) >= 20:
                retry = match_chunk(cr.chunk, blocks, set(all_idx), 1)
                if retry.hits:
                    cr.hits, cr.total_hits, cr.capped, cr.layer = \
                        retry.hits, retry.total_hits, True, retry.layer

    notes = []
    for k, cr in enumerate(cres):
        if not cr.hits:
            notes.append(f"第{k + 1}处摘录未定位（L1-L3 无命中）")
        if cr.capped:
            notes.append(f"第{k + 1}处命中 {cr.total_hits} 处仅标前 {len(cr.hits)} 处"
                         + ("" if filtered else "（定位区间未解析，仅挂首处）"))

    anchors = list(dict.fromkeys(i for cr in cres for i in cr.hits))
    status, l4 = "未匹配", None
    if anchors:
        status = "低置信" if any(cr.layer == "L3" for cr in cres if cr.hits) else "已批注"
    elif filtered:
        # L4 章节降级锚：挂过滤区间内首个标题块
        sec = next((s for s in sections
                    if s.start in cand and blocks[s.start].heading), None)
        if sec is not None:
            l4 = sec.start
            anchors = [l4]
            status = "章节降级"
            notes.append("摘录未定位，批注挂在本章节标题处（章节降级锚）")
    return FindingResult(row["F-ID"], row, cres, anchors, status, notes)


# ---------------------------------------------------------------- C 批注写入

def fmt_text(text: str) -> str:
    """台账文本 → 批注纯文本格式归一化：去 markdown 残留，§x.x → 第x.x节。"""
    t = (text or "").replace("\\|", "|")
    t = re.sub(r"§\s*([\d.]+)", r"第\1节", t)
    t = t.replace("**", "").replace("`", "")
    return t.strip()


def _release_base(val: str) -> str:
    m = re.match(r"^(阻断|条件放行|待澄清|不影响放行)", val or "")
    return m.group(1) if m else (val or "?")


def build_comment_text(row: dict, seq_total: int, seq_i: int,
                       table_rows: int | None, low_conf: bool) -> str:
    lines = []
    if low_conf:
        lines.append("⚠ 位置为模糊匹配，请人工核对")
    lines.append(f"【评审意见 {row['F-ID']}】{row.get('impact_level', '?')} · "
                 f"{_release_base(row.get('release_effect', ''))}（{row.get('evidence_status', '?')}）")
    lines.append(f"问题：{fmt_text(row.get('问题描述与后果', ''))}")
    if (row.get("修改建议") or "").strip():
        lines.append(f"建议：{fmt_text(row['修改建议'])}")
    if row.get("impact_level") in ("P0", "P1") and (row.get("关闭条件") or "").strip():
        lines.append(f"关闭条件：{fmt_text(row['关闭条件'])}")
    q = (row.get("澄清问题") or "").strip()
    if q and q != "无":
        lines.append(f"待澄清：{fmt_text(q)}")
    refs = set(re.findall(r"\bF-\d{3,}\b",
                          (row.get("问题描述与后果", "") or "")
                          + (row.get("修改建议", "") or "")
                          + (row.get("定位", "") or "")))
    refs.discard(row.get("F-ID", ""))
    if refs:
        lines.append("关联：" + "、".join(sorted(refs)))
    if table_rows:
        lines.append(f"本问题涉及本表 {table_rows} 行")
    elif seq_total > 1:
        lines.append(f"本问题共 {seq_total} 处标注（本处 {seq_i}/{seq_total}）")
    lines.append(SIGNATURE)
    return "\n".join(lines)


def plan_comment_points(anchors, blocks):
    """表格合并：同表 ≥2 锚合并为首个锚一条批注。返回 [(block_idx, table_rows|None)]。"""
    by_table = {}
    for i in anchors:
        b = blocks[i]
        if b.kind == "trow":
            by_table.setdefault(b.table_id, []).append(i)
    points, done_tables = [], set()
    for i in anchors:
        b = blocks[i]
        if b.kind == "trow":
            rows = by_table[b.table_id]
            if len(rows) >= 2:
                if b.table_id in done_tables:
                    continue
                done_tables.add(b.table_id)
                points.append((i, len(rows)))
                continue
        points.append((i, None))
    return points


def write_comments(doc, results: list, blocks) -> int:
    """写入全部批注，回填各 Finding 的实际写入处数，返回总写入条数。

    fail-closed：命中块全部不可锚导致 placed=0 时，状态强制归"未匹配"——
    状态与实际写入必须一致（防"已批注"却无批注的失实状态）。"""
    written = 0
    for fr in results:
        if fr.status == "未匹配" or not fr.anchors:
            continue
        low_conf = fr.status == "低置信"
        points = plan_comment_points(fr.anchors, blocks)
        for i, (bidx, table_rows) in enumerate(points, 1):
            b = blocks[bidx]
            if b.anchor is None:
                fr.notes.append(f"块{bidx}不可锚（空段落/零 run），已跳过")
                continue
            runs = _all_runs(b.anchor)
            if not runs:
                fr.notes.append(f"块{bidx}不可锚（空段落/零 run），已跳过")
                continue
            text = build_comment_text(fr.row, len(points), i, table_rows, low_conf)
            doc.add_comment(runs=runs, text=text, author=AUTHOR, initials=INITIALS)
            written += 1
            fr.placed += 1
        if fr.placed == 0:
            fr.status = "未匹配"
            fr.notes.append("命中块均不可锚，状态由降级为未匹配（fail-closed：不得虚报已批注）")
    return written


# ---------------------------------------------------------------- D 状态文件与验证

def count_rr_comments(path: str) -> int:
    """zipfile+正则数副本中 author=本工具 的批注数（与 validator 同口径，零 python-docx 依赖）。"""
    try:
        with zipfile.ZipFile(path) as zf:
            if "word/comments.xml" not in zf.namelist():
                return 0
            xml = zf.read("word/comments.xml").decode("utf-8", errors="ignore")
    except (OSError, zipfile.BadZipFile):
        return 0
    return sum(1 for _id, a in COMMENT_COUNT_PAT.findall(xml) if a == AUTHOR)


def sha256_of(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for blk in iter(lambda: f.read(1 << 20), b""):
            h.update(blk)
    return h.hexdigest()


def write_status_file(workdir, results, src_name, out_name, written, allow_unmatched):
    marked = [fr for fr in results if fr.status != "未匹配"]
    n_low = sum(1 for fr in results if fr.status == "低置信")
    lines = [
        "# docx 回写状态（writeback_docx.py 自动生成）",
        "",
        f"- 原件: {src_name}",
        f"- 副本: {out_name}",
        f"- 生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"- 覆盖率: {len(marked)}/{len(results)}"
        + (f"（低置信 {n_low}）" if n_low else ""),
        f"- 批注写入数: {written}（author={AUTHOR}）",
        f"- 未匹配豁免: {'是（--allow-unmatched）' if allow_unmatched else '无'}",
        "",
        "| F-ID | 片段数 | 已标注处数 | 状态 | 备注 |",
        "|---|---|---|---|---|",
    ]
    for fr in results:
        note = "；".join(fr.notes) if fr.notes else ""
        lines.append(f"| {fr.fid} | {len(fr.chunks)} | {fr.placed} | {fr.status} | {note} |")
    unmatched = [fr for fr in results if fr.status == "未匹配"]
    if unmatched:
        lines += ["", "## 未匹配明细", ""]
        for fr in unmatched:
            for k, cr in enumerate(fr.chunks, 1):
                lines.append(f"- {fr.fid} 第{k}处摘录「{cr.chunk}」：L1/L1.5/L2/L2.5/L3 均未命中"
                             "（摘录非原文或属转换噪声，请人工核对台账与原文）")
    p = os.path.join(workdir, "docx-writeback.md")
    with open(p, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


def find_source(workdir: str):
    """定位原件 source*.docx。排除本工具自己生成的批注副本（其文件名同样以 source 开头，
    不排除会在二次运行时把副本当原件、批注越叠越多）。validator section 11 的触发判定须用同一排除口径。"""
    cands = sorted(x for x in os.listdir(workdir)
                   if x.lower().endswith(".docx") and x.startswith("source")
                   and "评审批注" not in x)
    if not cands:
        return None
    return os.path.join(workdir, cands[0])


def run(workdir: str, allow_unmatched=False, cap=DEFAULT_CAP) -> int:
    src = find_source(workdir)
    if not src:
        print("错误: 工作目录内未找到 source*.docx", file=sys.stderr)
        return 2
    ledger = os.path.join(workdir, "issues-ledger.md")
    if not os.path.exists(ledger):
        print("错误: 缺少 issues-ledger.md", file=sys.stderr)
        return 2
    cols, rows, malformed = parse_ledger(ledger)
    if cols is None:
        print("错误: issues-ledger.md 表头缺失", file=sys.stderr)
        return 2
    if malformed:
        print("错误: 台账存在畸形行，先过 validate_review.py：", file=sys.stderr)
        for m in malformed:
            print(f"  - {m}", file=sys.stderr)
        return 2

    open_rows = [r for r in rows
                 if r.get("F-ID") and not r["F-ID"].startswith("_")
                 and r.get("状态") == "打开"]
    if not open_rows:
        print("警告: 台账中无 状态=打开 的 Finding，产出空状态文件")

    src_sha_before = sha256_of(src)
    doc = Document(src)
    blocks = build_blocks(doc)
    sections = build_sections(blocks)

    results = [match_finding(r, blocks, sections, cap) for r in open_rows]
    written = write_comments(doc, results, blocks)

    today = datetime.date.today().strftime("%Y%m%d")
    src_name = os.path.basename(src)
    out_name = f"{os.path.splitext(src_name)[0]}-评审批注-{today}.docx"
    out_path = os.path.join(workdir, out_name)
    doc.save(out_path)

    # 机器侧三重验证：原件只读 / 批注计数（author 过滤）/ 状态自洽
    src_sha_after = sha256_of(src)
    if src_sha_before != src_sha_after:
        print("致命错误: 原件 sha256 发生变化（违反只读铁律）", file=sys.stderr)
        return 2
    counted = count_rr_comments(out_path)
    if counted != written:
        print(f"致命错误: 副本批注计数 {counted} ≠ 写入数 {written}", file=sys.stderr)
        return 2

    write_status_file(workdir, results, src_name, out_name, written, allow_unmatched)

    n_unmatched = sum(1 for fr in results if fr.status == "未匹配")
    n_marked = len(results) - n_unmatched
    cov = f"{n_marked}/{len(results)}" if results else "0/0"
    print(json.dumps({
        "source": src_name, "output": out_name, "status_file": "docx-writeback.md",
        "coverage": cov, "written_comments": written, "unmatched": n_unmatched,
        "by_status": {s: sum(1 for fr in results if fr.status == s) for s in VALID_WB_STATUS},
    }, ensure_ascii=False))
    if n_unmatched and not allow_unmatched:
        return 1
    return 0


def main():
    ap = argparse.ArgumentParser(description="评审台账回写 docx 批注副本")
    ap.add_argument("workdir", help="评审工作目录（含 source*.docx 与 issues-ledger.md）")
    ap.add_argument("--allow-unmatched", action="store_true",
                    help="豁免未匹配 Finding（须在报告中说明原因）")
    ap.add_argument("--cap", type=int, default=DEFAULT_CAP,
                    help=f"单摘录多处命中的挂注上限（默认 {DEFAULT_CAP}）")
    args = ap.parse_args()
    sys.exit(run(args.workdir, args.allow_unmatched, args.cap))


if __name__ == "__main__":
    main()
